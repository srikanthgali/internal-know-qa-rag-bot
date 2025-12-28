import os
import logging
from pathlib import Path
from typing import List, Dict, Optional, Union, Any
from dataclasses import dataclass
from datetime import datetime
from src.utils.logger import get_logger

# Document processing libraries
import PyPDF2
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter


logger = get_logger(__name__)


@dataclass
class Document:
    """Represents a document chunk with metadata."""

    content: str
    metadata: Dict[str, Any]  # Changed from 'any' to 'Any'
    doc_id: Optional[str] = None


class DocumentLoader:
    """Loads and processes documents from various file formats."""

    SUPPORTED_FORMATS = {".pdf", ".txt", ".docx", ".md"}

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: Optional[List[str]] = None,
    ):
        """
        Initialize the document loader.

        Args:
            chunk_size: Maximum size of each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
            separators: Custom separators for text splitting
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        if separators is None:
            separators = ["\n\n", "\n", ". ", " ", ""]

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
            length_function=len,
        )

    def load_document(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a document from file path.

        Args:
            file_path: Path to the document file

        Returns:
            List of Document objects with chunked content
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_path.suffix.lower() not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {file_path.suffix}. "
                f"Supported formats: {self.SUPPORTED_FORMATS}"
            )

        logger.info(f"Loading document: {file_path}")

        # Extract text based on file type
        if file_path.suffix.lower() == ".pdf":
            text = self._load_pdf(file_path)
        elif file_path.suffix.lower() == ".docx":
            text = self._load_docx(file_path)
        elif file_path.suffix.lower() in {".txt", ".md"}:
            text = self._load_text(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_path.suffix}")

        # Validate extracted text
        if not text or not text.strip():
            logger.warning(f"No text content extracted from {file_path}")
            return []

        # Create base metadata
        metadata = self._extract_metadata(file_path)

        # Split into chunks
        chunks = self.text_splitter.split_text(text)

        # Create Document objects
        documents = []
        for idx, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update(
                {
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                    "chunk_size": len(chunk),
                }
            )

            doc = Document(
                content=chunk, metadata=chunk_metadata, doc_id=f"{file_path.stem}_{idx}"
            )
            documents.append(doc)

        logger.info(f"Loaded {len(documents)} chunks from {file_path.name}")
        return documents

    def load_directory(
        self, directory_path: Union[str, Path], recursive: bool = True
    ) -> List[Document]:
        """
        Load all supported documents from a directory.

        Args:
            directory_path: Path to the directory
            recursive: Whether to search subdirectories

        Returns:
            List of all Document objects from the directory
        """
        directory_path = Path(directory_path)

        if not directory_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory_path}")

        all_documents = []

        # Get all files
        if recursive:
            files = list(directory_path.rglob("*"))
        else:
            files = list(directory_path.glob("*"))

        # Filter for supported formats
        supported_files = [
            f
            for f in files
            if f.is_file() and f.suffix.lower() in self.SUPPORTED_FORMATS
        ]

        logger.info(f"Found {len(supported_files)} supported files in {directory_path}")

        # Load each file
        for file_path in supported_files:
            try:
                documents = self.load_document(file_path)
                all_documents.extend(documents)
            except Exception as e:
                logger.error(f"Error loading {file_path}: {str(e)}")
                continue

        logger.info(f"Loaded total of {len(all_documents)} document chunks")
        return all_documents

    def _load_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text_content = []

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                    except Exception as e:
                        logger.warning(
                            f"Error extracting page {page_num} from {file_path}: {e}"
                        )
                        continue

        except Exception as e:
            raise RuntimeError(f"Error reading PDF file {file_path}: {str(e)}")

        return "\n\n".join(text_content)

    def _load_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise RuntimeError(f"Error reading DOCX file {file_path}: {str(e)}")

    def _load_text(self, file_path: Path) -> str:
        """Load plain text or markdown file."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as file:
                return file.read()
        except Exception as e:
            raise RuntimeError(f"Error reading text file {file_path}: {str(e)}")

    def _extract_metadata(
        self, file_path: Path
    ) -> Dict[str, Any]:  # Changed from 'any' to 'Any'
        """Extract metadata from file."""
        stat = file_path.stat()

        return {
            "source": str(file_path),
            "filename": file_path.name,
            "file_type": file_path.suffix.lower(),
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        }


# Example usage
if __name__ == "__main__":
    # Initialize loader
    loader = DocumentLoader(chunk_size=1000, chunk_overlap=200)

    # Example 1: Load single document
    try:
        documents = loader.load_document("path/to/document.pdf")
        print(f"Loaded {len(documents)} chunks")

        # Print first chunk
        if documents:
            print(f"\nFirst chunk preview:")
            print(f"Content: {documents[0].content[:200]}...")
            print(f"Metadata: {documents[0].metadata}")
    except Exception as e:
        print(f"Error: {e}")

    # Example 2: Load directory
    try:
        documents = loader.load_directory("path/to/documents", recursive=True)
        print(f"\nLoaded {len(documents)} total chunks from directory")
    except Exception as e:
        print(f"Error: {e}")
